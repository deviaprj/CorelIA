#!/usr/bin/env python3
"""Convertit docs/notice_api_keys.md en DOCX avec mise en forme professionnelle."""

import re
import sys
from pathlib import Path

try:
    from docx import Document
    import docx.opc.constants
    import docx.oxml
    from docx.shared import Inches, Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
except ImportError:
    print("python-docx non installé. Installez-le via : pip install python-docx")
    sys.exit(1)


def add_hyperlink(paragraph, text, url):
    """Ajoute un lien cliquable dans un paragraphe."""
    part = paragraph.part
    r_id = part.relate_to(
        url,
        docx.opc.constants.RELATIONSHIP_TYPE.HYPERLINK,
        is_external=True
    )
    hyperlink = docx.oxml.parse_xml(
        r'<w:hyperlink xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main" r:id="%s" xmlns:r="http://schemas.openxmlformats.org/officeDocument/2006/relationships">'
        % r_id
    )
    new_run = docx.oxml.parse_xml(
        r'<w:r xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">'
        r'<w:rPr><w:rStyle w:val="Hyperlink"/><w:color w:val="6C63FF"/><w:u w:val="single"/></w:rPr>'
        r'<w:t>%s</w:t></w:r>' % text
    )
    hyperlink.append(new_run)
    paragraph._p.append(hyperlink)
    return hyperlink


def format_code_block(doc, lines, start_idx):
    """Ajoute un bloc de code stylisé."""
    i = start_idx
    code_lines = []
    while i < len(lines) and not lines[i].strip().startswith('```'):
        code_lines.append(lines[i])
        i += 1
    if i < len(lines):
        i += 1

    if code_lines:
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Inches(0.3)
        p.paragraph_format.space_before = Pt(6)
        p.paragraph_format.space_after = Pt(6)
        run = p.add_run('\n'.join(code_lines))
        run.font.name = 'Courier New'
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(0x33, 0x33, 0x33)
    return i


def format_table(doc, lines, start_idx):
    """Ajoute un tableau Markdown stylisé."""
    i = start_idx
    header = [c.strip() for c in lines[i].split('|') if c.strip()]
    i += 2  # sauter la ligne séparateur |---|---|

    rows_data = []
    while i < len(lines) and lines[i].strip().startswith('|'):
        cells = [c.strip() for c in lines[i].split('|') if c.strip()]
        if cells:
            rows_data.append(cells)
        i += 1

    if header and rows_data:
        table = doc.add_table(rows=1 + len(rows_data), cols=len(header))
        table.style = 'Table Grid'
        # Header
        for j, h in enumerate(header):
            cell = table.rows[0].cells[j]
            cell.text = h
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.bold = True
                    run.font.size = Pt(10)
                    run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            shading = docx.oxml.parse_xml(
                r'<w:shd xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main" w:fill="6C63FF" w:val="clear"/>'
            )
            cell._tc.get_or_add_tcPr().append(shading)

        # Rows
        for row_idx, row_data in enumerate(rows_data):
            for j, cell_text in enumerate(row_data):
                if j < len(header):
                    cell = table.rows[row_idx + 1].cells[j]
                    cell.text = cell_text
                    for paragraph in cell.paragraphs:
                        for run in paragraph.runs:
                            run.font.size = Pt(9)
        doc.add_paragraph()
    return i


def convert_md_to_docx(md_path, docx_path):
    doc = Document()

    # Styles de base
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Inter'
    font.size = Pt(11)
    font.color.rgb = RGBColor(0x22, 0x22, 0x22)

    md_content = Path(md_path).read_text(encoding='utf-8')
    lines = md_content.splitlines()

    i = 0
    while i < len(lines):
        line = lines[i]
        stripped = line.strip()

        if stripped == '---':
            doc.add_page_break()
            i += 1
            continue

        # Heading 1
        if stripped.startswith('# ') and not stripped.startswith('## '):
            p = doc.add_heading(stripped[2:], level=1)
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            for run in p.runs:
                run.font.color.rgb = RGBColor(0x6C, 0x63, 0xFF)
                run.font.size = Pt(20)
            i += 1
            continue

        # Heading 2
        if stripped.startswith('## '):
            p = doc.add_heading(stripped[3:], level=2)
            for run in p.runs:
                run.font.color.rgb = RGBColor(0x6C, 0x63, 0xFF)
                run.font.size = Pt(16)
            i += 1
            continue

        # Heading 3
        if stripped.startswith('### '):
            p = doc.add_heading(stripped[4:], level=3)
            for run in p.runs:
                run.font.color.rgb = RGBColor(0x44, 0x44, 0x44)
                run.font.size = Pt(13)
                run.font.bold = True
            i += 1
            continue

        # Heading 4
        if stripped.startswith('#### '):
            p = doc.add_heading(stripped[5:], level=4)
            for run in p.runs:
                run.font.color.rgb = RGBColor(0x44, 0x44, 0x44)
                run.font.size = Pt(12)
                run.font.bold = True
            i += 1
            continue

        # Bloc de code
        if stripped.startswith('```'):
            i = format_code_block(doc, lines, i + 1)
            continue

        # Tableau
        if stripped.startswith('|') and i + 1 < len(lines) and '---' in lines[i + 1]:
            i = format_table(doc, lines, i)
            continue

        # Ligne vide
        if not stripped:
            doc.add_paragraph()
            i += 1
            continue

        # Blockquote
        if stripped.startswith('> '):
            p = doc.add_paragraph(stripped[2:])
            p.paragraph_format.left_indent = Inches(0.2)
            for run in p.runs:
                run.font.italic = True
                run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
            i += 1
            continue

        # Liste à puces
        if re.match(r'^(\-|\*|\[ \]|\[x\])\s+', stripped):
            text = re.sub(r'^(\-|\*|\[ \]|\[x\])\s+', '', stripped)
            p = doc.add_paragraph(text, style='List Bullet')
            p.paragraph_format.left_indent = Inches(0.25)
            i += 1
            continue

        # Liste numérotée
        if re.match(r'^\d+\.\s+', stripped):
            text = re.sub(r'^\d+\.\s+', '', stripped)
            p = doc.add_paragraph(text, style='List Number')
            i += 1
            continue

        # Paragraphe normal avec formatage inline
        p = doc.add_paragraph()
        parts = re.split(r'(\*\*.*?\*\*|\*.*?\*|`[^`]+?`|\[.*?\]\(.*?\))', stripped)
        for part in parts:
            if part.startswith('**') and part.endswith('**'):
                run = p.add_run(part[2:-2])
                run.bold = True
            elif part.startswith('*') and part.endswith('*') and not part.startswith('**'):
                run = p.add_run(part[1:-1])
                run.italic = True
            elif part.startswith('`') and part.endswith('`'):
                run = p.add_run(part[1:-1])
                run.font.name = 'Courier New'
                run.font.size = Pt(10)
                run.font.color.rgb = RGBColor(0xCC, 0x33, 0x33)
            elif re.match(r'\[([^\]]+)\]\(([^)]+)\)', part):
                m = re.match(r'\[([^\]]+)\]\(([^)]+)\)', part)
                add_hyperlink(p, m.group(1), m.group(2))
            else:
                p.add_run(part)
        i += 1

    # Marges
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)

    # Footer
    footer = section.footer
    footer_para = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
    footer_para.text = "Corely — Notice API Keys — Document confidentiel"
    footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in footer_para.runs:
        run.font.size = Pt(8)
        run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)

    doc.save(docx_path)
    print(f"✅ Document généré : {docx_path}")


if __name__ == '__main__':
    root = Path(__file__).parent.parent
    md_file = root / 'docs' / 'notice_api_keys.md'
    docx_file = root / 'docs' / 'notice_api_keys.docx'
    convert_md_to_docx(str(md_file), str(docx_file))
